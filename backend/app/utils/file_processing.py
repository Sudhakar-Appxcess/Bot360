from fastapi import UploadFile
from PyPDF2 import PdfReader
import pandas as pd
from docx import Document
from PIL import Image
import pytesseract
import io

async def extract_text_from_file(file: UploadFile) -> str:
    content = ""
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        content = await extract_text_from_pdf(file)
    elif filename.endswith(".csv"):
        content = await extract_text_from_csv(file)
    elif filename.endswith(".docx"):
        content = await extract_text_from_word(file)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        content = await extract_text_from_image(file)
    
    return content

async def extract_text_from_pdf(file: UploadFile) -> str:
    reader = PdfReader(file.file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

async def extract_text_from_csv(file: UploadFile) -> str:
    print("CSV")
    df = pd.read_csv(file.file)
    return df.to_string(index=False)

# async def extract_text_from_word(file: UploadFile) -> str:
#     print("WORD")
#     doc = Document(file.file)
#     text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#     return text
from docx import Document
from io import BytesIO

async def extract_text_from_word(file):
    file_content = await file.read()
    docx_file = BytesIO(file_content)
    doc = Document(docx_file)
    content = "\n".join([para.text for para in doc.paragraphs])
    return content


async def extract_text_from_image(file: UploadFile) -> str:
    image = Image.open(io.BytesIO(await file.read()))
    text = pytesseract.image_to_string(image)
    return text
